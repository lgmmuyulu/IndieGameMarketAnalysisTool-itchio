# -*- coding: utf-8 -*-
import pandas as pd
import numpy as np
import re
import os
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# --- 字体配置 ---
def setup_chinese_font():
    """查找并设置可用的中文字体"""
    font_list = ['Microsoft YaHei', 'SimHei', 'PingFang SC', 'STHeiti', 'WenQuanYi Micro Hei']
    import matplotlib.font_manager as fm
    for font in font_list:
        if font in [f.name for f in fm.fontManager.ttflist]:
            plt.rcParams['font.sans-serif'] = [font]
            print(f"Font set to: {font}")
            return
    print("Warning: No recommended Chinese font found. Chinese characters may not display correctly.")

setup_chinese_font()
plt.rcParams['axes.unicode_minus'] = False

class GameReviewAnalyzer:
    """
    从CSV文件加载游戏评论数据，执行分析并生成图文并茂的Word报告。
    """
    
    def __init__(self, game_name):
        self.game_name = game_name
        self.reviews_df = None
        self.analysis_result = {}

    def load_data(self, file_path):
        """
        加载CSV格式的评论数据。
        必需列: review_text, rating
        """
        try:
            self.reviews_df = pd.read_csv(file_path, encoding='utf-8')
            self.reviews_df = self.reviews_df.dropna(subset=['review_text', 'rating'])
            self.reviews_df['rating'] = pd.to_numeric(self.reviews_df['rating'], errors='coerce')
            self.reviews_df = self.reviews_df.dropna(subset=['rating'])
            self.reviews_df['rating_10_scale'] = self.reviews_df['rating'] * 2
            print(f"Loaded {len(self.reviews_df)} reviews from {file_path}")
            return self
        except FileNotFoundError:
            print(f"Warning: '{file_path}' not found. Using mock data for demonstration.")
            self.reviews_df = self._generate_mock_data()
            return self
        except Exception as e:
            print(f"Error loading data: {e}")
            return None

    def sentiment_analysis(self):
        """根据评分进行情感分类"""
        def classify(rating_10):
            if rating_10 >= 9: return '极度好评'
            elif rating_10 >= 7: return '好评'
            elif rating_10 >= 5: return '中性'
            elif rating_10 >= 3: return '差评'
            else: return '极度差评'
        
        self.reviews_df['sentiment'] = self.reviews_df['rating_10_scale'].apply(classify)
        self.analysis_result['sentiment_distribution'] = self.reviews_df['sentiment'].value_counts()
        return self

    def dimension_scoring(self):
        """计算六个维度的得分 (究极版关键词词典)"""
        keywords = {
            '核心玩法': [
                # --- General Gameplay ---
                '玩法', '游戏性', '可玩性', '核心', '机制', '系统', '耐玩', '好玩', '有趣', '上头', '沉迷', '上瘾', '快乐', '有意思', '享受',
                'gameplay', 'fun', 'addictive', 'mechanics', 'core', 'system', 'loop', 'playability', 'enjoyable', 'engaging', 'compelling',
                '无聊', '枯燥', '重复', '单调', '没意思', '不好玩', '乏味', '催眠', 'boring', 'repetitive', 'dull', 'tedious', 'shallow', 'uninspired', 'grindy',
                # --- Combat ---
                '战斗', '打击感', '战斗系统', '动作', '技能', '武器', '敌人', 'boss', '战斗体验', '爽', '连招',
                'combat', 'action', 'skills', 'weapons', 'enemies', 'bosses', 'fighting', 'feel', 'impact', 'satisfying',
                # --- Puzzle ---
                '解谜', '谜题', '关卡', '设计', '巧妙', '烧脑', '智商', 'puzzles', 'riddles', 'level design', 'clever', 'challenging', 'brainy',
                # --- Difficulty & Progression ---
                '难度', '挑战', '硬核', '简单', '容易', '太难', '劝退', 'difficulty', 'challenge', 'hard', 'easy', 'grind', 'punishing',
                '成长', '升级', '加点', '天赋', '装备', '养成', '构筑', 'BD', 'progression', 'leveling', 'upgrades', 'gear', 'talents', 'build',
                # --- Replayability ---
                '重玩价值', '多周目', '耐玩性', '内容', '时长', '太短', '流程', 'replayability', 'content', 'length', 'short', 'long', 'value'
            ],
            '叙事逻辑': [
                # --- General Narrative ---
                '剧情', '故事', '叙事', '情节', '主题', '故事线', '剧本', '文案', '文本', '感人', '深刻', '引人入胜', '扣人心弦', '发人深省',
                'story', 'narrative', 'plot', 'theme', 'writing', 'script', 'text', 'emotional', 'touching', 'compelling', 'engaging', 'moving', 'thought-provoking',
                # --- Characters ---
                '角色', '人物', '主角', '配角', '塑造', '立体', '丰满', '有魅力', 'character', 'protagonist', 'cast', 'development', 'relatable', 'memorable',
                # --- World & Lore ---
                '世界观', '设定', '背景', '代入感', '沉浸感', 'lore', 'world-building', 'setting', 'immersion', 'immersive', 'atmosphere',
                # --- Dialogue & Pacing ---
                '对话', '配音', 'dialogue', 'voice acting', 'va', 'voiceover',
                '节奏', ' pacing', 'ending', '结局', '高潮', 'climax', 'twist', '反转', '伏笔', 'foreshadowing'
            ],
            '美术风格': [
                # --- General Art & Style ---
                '美术', '画面', '画风', '风格', '视觉', '好看', '精美', '惊艳', '艺术', '设计', '审美', '漂亮', '华丽', '舒服',
                'art', 'visuals', 'beautiful', 'style', 'design', 'gorgeous', 'stunning', 'aesthetics', 'look', 'pretty', 'lovely',
                # --- Graphics & Models ---
                '建模', '模型', '贴图', '材质', '渲染', '光影', '特效', '粒子', '画质',
                'graphics', 'models', 'textures', 'rendering', 'lighting', 'effects', 'vfx', 'fidelity',
                # --- Environment & Animation ---
                '场景', '环境', '地图', '动画', '动作', '流畅', '僵硬', '细节',
                'environment', 'scene', 'map', 'animation', 'smooth', 'janky', 'stiff', 'fluid', 'detail',
                # --- UI ---
                'UI', '界面', '用户界面', '菜单', '图标', '简洁', '清晰', 'interface', 'hud', 'menu', 'icons', 'clean', 'minimalist'
            ],
            '音效配乐': [
                # --- Music ---
                '音乐', '配乐', '原声', 'BGM', '主题曲', '好听', '神曲', '旋律',
                'music', 'soundtrack', 'score', 'ost', 'bgm', 'theme', 'melody', 'tune',
                # --- Sound Design ---
                '音效', '声音', '环境音', '打击音', '反馈', '细节',
                'sound', 'audio', 'sfx', 'sound design', 'feedback', 'details',
                # --- Voice & Atmosphere ---
                '配音', '声优', '语音', 'CV', 'voice acting', 'vo', 'voiceover', 'narration',
                '氛围', '气氛', 'ambience', 'atmosphere'
            ],
            '用户体验': [
                # --- Bugs & Performance ---
                'bug', 'BUG', '漏洞', '问题', '错误', '闪退', '崩溃', '死机', '坏档', '恶性',
                'bugs', 'glitches', 'issues', 'errors', 'crash', 'freeze', 'corrupted save', 'showstopper',
                '优化', '性能', '卡顿', '掉帧', '帧率', '加载', '延迟', '拖慢', '撕裂',
                'optimization', 'performance', 'lag', 'stutter', 'fps', 'frame rate', 'loading', 'delay', 'tearing',
                # --- Controls & UI/UX ---
                '操作', '手感', '按键', '控制', '摇杆', '鼠标', '键盘', '响应', '别扭', '反人类',
                'controls', 'handling', 'input', 'controller', 'mouse', 'keyboard', 'responsive', 'clunky',
                '引导', '教程', '新手', '提示', '教学', '指引', 'tutorial', 'guidance', 'onboarding',
                'UI', 'UX', '界面', '交互', '菜单', '方便', '直观', '易用',
                'interface', 'user experience', 'menu', 'intuitive', 'user-friendly'
            ],
            '创新性': [
                # --- Originality & Creativity ---
                '创新', '创意', '新颖', '独特', '原创', '点子', '想法', '脑洞', '首创', '开创',
                'innovative', 'creative', 'unique', 'original', 'idea', 'concept', 'groundbreaking',
                # --- Freshness ---
                '新鲜', '新意', '特别', '不一样', '惊喜', '眼前一亮',
                'fresh', 'new', 'different', 'special', 'surprise', 'refreshing',
                # --- Genre ---
                '类型', '融合', '玩法融合', '结合', 'genre', 'mix', 'hybrid', 'blend'
            ]
        }
        scores = {}
        for dim, kws in keywords.items():
            pattern = '|'.join(kws)
            relevant_df = self.reviews_df[self.reviews_df['review_text'].str.contains(pattern, case=False, na=False)]
            if not relevant_df.empty:
                avg_score = relevant_df['rating_10_scale'].mean()
                count = len(relevant_df)
            else:
                avg_score = self.reviews_df['rating_10_scale'].mean()
                count = 0
            scores[dim] = {'score': round(avg_score, 1), 'mentions': count}
        self.analysis_result['dimension_scores'] = scores
        return self

    def generate_word_report(self, output_path='analysis_report.docx'):
        """将分析结果和图表写入Word文档"""
        doc = Document()
        doc.add_heading(f'{self.game_name} - 玩家评论分析报告', level=0)
        
        doc.add_heading('执行摘要', level=1)
        doc.add_paragraph(f"总评论数: {len(self.reviews_df)}", style='List Bullet')
        doc.add_paragraph(f"综合评分 (10分制): {self.reviews_df['rating_10_scale'].mean():.1f}/10", style='List Bullet')
        doc.add_paragraph(f"好评率 (4-5星): {(self.reviews_df['rating'] >= 4).sum() / len(self.reviews_df) * 100:.1f}%", style='List Bullet')

        doc.add_heading('六维评估', level=1)
        try:
            doc.add_picture('./charts/radar_chart.png', width=Inches(6.0))
        except FileNotFoundError:
            doc.add_paragraph("雷达图文件未找到。")
        for dim, data in self.analysis_result.get('dimension_scores', {}).items():
            doc.add_paragraph(f"{dim}: {data['score']}/10 (提及 {data['mentions']} 次)", style='List Bullet 2')
        
        doc.add_heading('情感分布', level=1)
        try:
            doc.add_picture('./charts/sentiment_pie.png', width=Inches(5.0))
        except FileNotFoundError:
            doc.add_paragraph("饼图文件未找到。")
        for sentiment, count in self.analysis_result.get('sentiment_distribution', {}).items():
            pct = count / len(self.reviews_df) * 100
            doc.add_paragraph(f"{sentiment}: {count}条 ({pct:.1f}%)", style='List Bullet 2')

        doc.save(output_path)
        print(f"Report saved to {output_path}")
        return self

    def visualize(self, output_dir='./charts/'):
        """生成并保存分析图表"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 情感分布饼图
        sentiment_data = self.analysis_result['sentiment_distribution']
        plt.figure(figsize=(8, 6))
        plt.pie(sentiment_data.values, labels=sentiment_data.index, autopct='%1.1f%%', startangle=90,
                colors=['#10b981', '#6ee7b7', '#fbbf24', '#f87171', '#dc2626'])
        plt.title(f'{self.game_name} - 情感分布')
        plt.savefig(f'{output_dir}sentiment_pie.png', dpi=300, bbox_inches='tight')
        plt.close()

        # 六维雷达图
        dim_scores = self.analysis_result['dimension_scores']
        labels = list(dim_scores.keys())
        stats = [d['score'] for d in dim_scores.values()]
        angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
        stats += stats[:1]
        angles += angles[:1]
        
        fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
        ax.plot(angles, stats, 'o-', linewidth=2)
        ax.fill(angles, stats, alpha=0.25)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(labels)
        ax.set_ylim(0, 10)
        ax.set_title(f'{self.game_name} - 六维评估', pad=20)
        plt.savefig(f'{output_dir}radar_chart.png', dpi=300, bbox_inches='tight')
        plt.close()

        print(f"Visualizations saved to {output_dir}")
        return self

    def run_full_analysis(self):
        """执行完整的分析流程"""
        if self.reviews_df is None:
            print("Data not loaded. Aborting analysis.")
            return None
        self.sentiment_analysis()
        self.dimension_scoring()
        print("Analysis complete.")
        return self
        
    def _generate_mock_data(self, n_reviews=200):
        """内部方法，用于生成模拟数据"""
        np.random.seed(42)
        pos_templates = ["玩法很有趣", "画面精美", "剧情优秀", "An incredibly unique concept!", "A beautiful little story."]
        neg_templates = ["有点重复", "优化不好", "bug较多", "Boring gameplay", "The controls are a bit clunky"]
        reviews = []
        for _ in range(n_reviews):
            rating = np.random.choice([1, 2, 3, 4, 5], p=[0.05, 0.05, 0.17, 0.16, 0.57])
            text = np.random.choice(pos_templates) if rating >= 4 else np.random.choice(neg_templates)
            reviews.append({'review_text': text, 'rating': rating})
        df = pd.DataFrame(reviews)
        df['rating_10_scale'] = df['rating'] * 2
        return df

def main():
    """主函数"""
    analyzer = GameReviewAnalyzer(game_name="Love Letters")
    analyzer = analyzer.load_data('itchio_reviews.csv')
    if analyzer:
        analyzer.run_full_analysis()
        analyzer.visualize()
        analyzer.generate_word_report()

if __name__ == "__main__":
    main()

